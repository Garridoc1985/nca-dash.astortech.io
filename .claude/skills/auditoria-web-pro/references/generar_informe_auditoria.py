"""
Generador de Informe de Auditoría Web — ASTOR
====================================================
Template genérico reutilizable. Para cada nueva auditoría:
  1. Copiar este archivo como: scripts/generar_informe_[dominio]-[fecha].py
  2. Reemplazar DATOS con los resultados reales del análisis
  3. Ajustar la ruta DESTINO al final
  4. Ejecutar: python scripts/generar_informe_[dominio]-[fecha].py

La función generar_informe(datos, ruta_destino) no requiere modificación.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx2pdf import convert
import os

# ─── PALETA ASTOR ─────────────────────────────────────────
OBSIDIANA  = RGBColor(0x1A, 0x1A, 0x18)
ORO        = RGBColor(0xC4, 0x96, 0x3A)
BLANCO     = RGBColor(0xFA, 0xF8, 0xF4)
PIEDRA     = RGBColor(0x6B, 0x68, 0x60)

def score_color(score):
    if score >= 85: return RGBColor(0x2D, 0x7A, 0x52)
    if score >= 75: return RGBColor(0x2D, 0x6B, 0x7A)
    if score >= 70: return RGBColor(0xC4, 0x96, 0x3A)
    if score >= 60: return RGBColor(0xA0, 0x78, 0x30)
    if score >= 40: return RGBColor(0xB0, 0x5C, 0x28)
    return RGBColor(0x92, 0x28, 0x28)

def score_label(score):
    if score >= 95: return "A+ · Excepcional"
    if score >= 90: return "A · Excelente"
    if score >= 85: return "B+ · Muy Bueno"
    if score >= 75: return "B · Bueno"
    if score >= 70: return "C+ · Aceptable"
    if score >= 60: return "C · Regular"
    if score >= 40: return "D · Necesita Mejoras"
    return "F · Crítico"

# ─── HELPERS ──────────────────────────────────────────────
def set_font(run, size_pt, bold=False, color=None, italic=False):
    run.font.name  = 'Arial'
    run.font.size  = Pt(size_pt)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def set_cell_bg(cell, color_hex):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  color_hex)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    set_font(run, {1: 14, 2: 12, 3: 11}.get(level, 11), bold=True, color=color or OBSIDIANA)
    return p

def add_body(doc, text, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.line_spacing = Pt(16)
    run = p.add_run(text)
    set_font(run, 11, color=color or PIEDRA)
    return p

def add_rule(doc):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr   = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'C4963A')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_finding(doc, icon, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(0.5)
    r1 = p.add_run(icon + '  ')
    set_font(r1, 10.5)
    r2 = p.add_run(text)
    set_font(r2, 10.5, color=PIEDRA)

def add_rec(doc, prioridad, titulo, detalle):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.line_spacing = Pt(15)
    r1 = p.add_run(f'[{prioridad}] ')
    set_font(r1, 10, bold=True, color=ORO)
    r2 = p.add_run(titulo + ' — ')
    set_font(r2, 10.5, bold=True, color=OBSIDIANA)
    r3 = p.add_run(detalle)
    set_font(r3, 10.5, color=PIEDRA)

def add_scores_table(doc, scores):
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, txt in enumerate(['Dimensión', 'Score', 'Calificación', 'Peso']):
        set_cell_bg(hdr[i], '1A1A18')
        p   = hdr[i].paragraphs[0]
        run = p.add_run(txt)
        set_font(run, 9, bold=True, color=BLANCO)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for dim, score, peso in scores:
        row = table.add_row().cells
        r0  = row[0].paragraphs[0].add_run(dim)
        set_font(r0, 10, color=OBSIDIANA)

        p1  = row[1].paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1  = p1.add_run(str(score))
        set_font(r1, 11, bold=True, color=score_color(score))

        r2  = row[2].paragraphs[0].add_run(score_label(score))
        set_font(r2, 9.5, color=PIEDRA)

        p3  = row[3].paragraphs[0]
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3  = p3.add_run(peso)
        set_font(r3, 9.5, color=PIEDRA)

    for i, row in enumerate(table.rows[1:]):
        bg = 'F0EDE6' if i % 2 == 0 else 'FFFFFF'
        for cell in row.cells:
            set_cell_bg(cell, bg)
    doc.add_paragraph()

def add_roadmap_table(doc, acciones):
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, txt in enumerate(['Acción', 'Dimensión', 'Prioridad', 'Impacto', 'Herramienta']):
        set_cell_bg(hdr[i], '1A1A18')
        p   = hdr[i].paragraphs[0]
        run = p.add_run(txt)
        set_font(run, 9, bold=True, color=BLANCO)

    for acc, dim, prio, imp, herr in acciones:
        row = table.add_row().cells
        for cell, txt, clr in [
            (row[0], acc,  OBSIDIANA),
            (row[1], dim,  PIEDRA),
            (row[2], prio, ORO if prio == 'Alta' else PIEDRA),
            (row[3], imp,  PIEDRA),
            (row[4], herr, PIEDRA),
        ]:
            run = cell.paragraphs[0].add_run(txt)
            set_font(run, 9.5, bold=(clr == ORO), color=clr)

    for i, row in enumerate(table.rows[1:]):
        bg = 'F0EDE6' if i % 2 == 0 else 'FFFFFF'
        for cell in row.cells:
            set_cell_bg(cell, bg)


# ─── ESTRUCTURA DE DATOS (completar por auditoría) ─────────
#
# DATOS = {
#     "sitio":        "nombre-del-sitio.cl",      # dominio sin https
#     "url":          "https://nombre-del-sitio.cl/",
#     "plataforma":   "WordPress · Elementor 3.x", # plataforma detectada
#     "fecha":        "DD de mes de AAAA",
#     "score_global": 72,                           # suma ponderada 0-100
#
#     "resumen": "Párrafo de 4-6 oraciones con los hallazgos principales y contexto ejecutivo.",
#
#     "top3": [
#         "Acción urgente 1 con tiempo estimado e impacto.",
#         "Acción urgente 2.",
#         "Acción urgente 3.",
#     ],
#
#     "scores": [
#         ("Performance & Carga",       75, "15%"),
#         ("SEO On-Page",               68, "15%"),
#         ("SEM & Tracking",            55, "10%"),
#         ("Diseño & Responsividad",    80, "15%"),
#         ("UX/UI & Accesibilidad",     65, "15%"),
#         ("Calidad del Contenido",     60, "10%"),
#         ("Seguridad Técnica",         70, "10%"),
#         ("Optimización Conversión",   72, "10%"),
#     ],
#
#     "dimensiones": [
#         {
#             "nombre":      "1. Performance & Carga — 75 / 100",
#             "score":       75,
#             "descripcion": "Párrafo de análisis contextual de la dimensión.",
#             "fortalezas":  ["Fortaleza 1.", "Fortaleza 2."],
#             "problemas":   ["Problema 1.", "Problema 2."],
#             "recs": [
#                 ("Alta",  "Título recomendación", "Detalle con pasos concretos."),
#                 ("Media", "Título recomendación", "Detalle con pasos concretos."),
#             ],
#         },
#         # ... repetir para las 8 dimensiones
#     ],
#
#     "roadmap": [
#         # (acción, dimensión, prioridad, impacto, herramienta)
#         ("Acción concreta 1", "SEO",        "Alta",  "Alto",  "Herramienta/plataforma"),
#         ("Acción concreta 2", "Performance","Media", "Medio", "Herramienta/plataforma"),
#     ],
# }


# ─── GENERADOR PRINCIPAL ──────────────────────────────────
def generar_informe(datos, ruta_destino):
    """
    Genera el .docx y lo convierte a PDF.
    ruta_destino: ruta completa del PDF (ej: r"C:/Users/Usuario/Desktop/Auditoria web Skills/auditoria-sitio-20260401.pdf")
    El .docx se guarda en la misma carpeta con el mismo nombre.
    """
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── PORTADA ──────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    run = p.add_run("ASTOR")
    set_font(run, 11, bold=True, color=ORO)
    run.font.all_caps = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run("Auditoría Web Integral")
    set_font(run, 22, bold=True, color=OBSIDIANA)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run(datos["sitio"])
    set_font(run, 16, italic=True, color=ORO)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(datos["url"])
    set_font(run, 10, color=PIEDRA)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(48)
    run = p.add_run(f"Plataforma: {datos['plataforma']}")
    set_font(run, 10, color=PIEDRA)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run(f"Fecha de auditoría: {datos['fecha']}")
    set_font(run, 10, color=PIEDRA)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    run = p.add_run(f"Puntaje Global: {datos['score_global']} / 100")
    set_font(run, 20, bold=True, color=score_color(datos['score_global']))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(score_label(datos['score_global']))
    set_font(run, 12, bold=True, color=score_color(datos['score_global']))

    doc.add_page_break()

    # ── RESUMEN EJECUTIVO ─────────────────────────────────
    add_heading(doc, "Resumen Ejecutivo", level=1)
    add_rule(doc)
    doc.add_paragraph()
    add_body(doc, datos["resumen"])

    add_heading(doc, "Top 3 Acciones Inmediatas", level=2)
    for i, rec in enumerate(datos["top3"], 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(0.5)
        r1 = p.add_run(f"{i}.  ")
        set_font(r1, 11, bold=True, color=ORO)
        r2 = p.add_run(rec)
        set_font(r2, 11, color=OBSIDIANA)

    doc.add_paragraph()

    # ── TABLA DE SCORES ───────────────────────────────────
    add_heading(doc, "Puntajes por Dimensión", level=1)
    add_rule(doc)
    doc.add_paragraph()
    add_scores_table(doc, datos["scores"])

    doc.add_page_break()

    # ── ANÁLISIS POR DIMENSIÓN ────────────────────────────
    add_heading(doc, "Análisis Detallado por Dimensión", level=1)
    add_rule(doc)

    for dim in datos["dimensiones"]:
        doc.add_paragraph()
        add_heading(doc, dim["nombre"], level=2, color=score_color(dim["score"]))
        add_body(doc, dim["descripcion"])

        add_heading(doc, "Hallazgos", level=3)
        for f in dim.get("fortalezas", []):
            add_finding(doc, "✅", f)
        for f in dim.get("problemas", []):
            add_finding(doc, "❌", f)

        add_heading(doc, "Recomendaciones", level=3)
        for prio, titulo, detalle in dim["recs"]:
            add_rec(doc, prio, titulo, detalle)

    doc.add_page_break()

    # ── ROADMAP ───────────────────────────────────────────
    add_heading(doc, "Roadmap Priorizado de Acciones", level=1)
    add_rule(doc)
    doc.add_paragraph()
    add_roadmap_table(doc, datos["roadmap"])

    doc.add_page_break()

    # ── METODOLOGÍA ───────────────────────────────────────
    add_heading(doc, "Notas Metodológicas", level=1)
    add_rule(doc)
    add_body(doc,
        f"Esta auditoría evalúa el estado observable del sitio {datos['url']} el {datos['fecha']}. "
        "El análisis cubre 8 dimensiones ponderadas: Performance (15%), SEO On-Page (15%), "
        "Diseño & Responsividad (15%), UX/UI & Accesibilidad (15%), SEM & Tracking (10%), "
        "Calidad del Contenido (10%), Seguridad Técnica (10%) y Optimización de Conversión (10%). "
        "El puntaje global es la suma ponderada de los 8 scores individuales."
    )
    add_body(doc,
        f"Contexto de plataforma: {datos['plataforma']}. "
        "La auditoría analiza lo observable desde el HTML estático de la homepage. "
        "No evalúa rendimiento de servidor (TTFB, Core Web Vitals reales) ni subpáginas. "
        "Complementar con Google PageSpeed Insights, GTmetrix y Screaming Frog para análisis completo."
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run("© 2026 ASTOR · astortech.io · Auditoría Digital · Informe confidencial")
    set_font(run, 9, color=PIEDRA)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── GUARDAR ───────────────────────────────────────────
    os.makedirs(os.path.dirname(ruta_destino), exist_ok=True)
    docx_path = ruta_destino.replace('.pdf', '.docx')
    doc.save(docx_path)
    print(f"DOCX generado: {docx_path} ({os.path.getsize(docx_path) // 1024} KB)")

    convert(docx_path, ruta_destino)
    print(f"PDF  generado: {ruta_destino} ({os.path.getsize(ruta_destino) // 1024} KB)")
    return docx_path, ruta_destino


# ─── MAIN ─────────────────────────────────────────────────
# Para usar este template en una nueva auditoría:
#   1. Copiar este archivo como scripts/generar_informe_[dominio]-[fecha].py
#   2. Definir DATOS = { ... } con los resultados reales
#   3. Actualizar DESTINO con el nombre del archivo de salida
#   4. Ejecutar: python scripts/generar_informe_[dominio]-[fecha].py
#
# Ejemplo mínimo:
#
# DATOS = {
#     "sitio": "ejemplo.cl", "url": "https://ejemplo.cl/",
#     "plataforma": "WordPress", "fecha": "01 de abril de 2026",
#     "score_global": 65,
#     "resumen": "...",
#     "top3": ["...", "...", "..."],
#     "scores": [("Performance & Carga", 70, "15%"), ...],
#     "dimensiones": [...],
#     "roadmap": [...]
# }
# DESTINO = r"C:/Users/Usuario/Desktop/Auditoria web Skills/auditoria-ejemplo-20260401.pdf"
# generar_informe(DATOS, DESTINO)

if __name__ == "__main__":
    print("Este es el template genérico. Copia este archivo, define DATOS y DESTINO, luego ejecuta.")
    print("Ver comentarios en la sección MAIN para instrucciones.")
